from __future__ import annotations

import io
import re
from email.message import EmailMessage
from email.policy import SMTP
from typing import Any, Dict, List, Tuple

import requests
from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
LINK_RE = re.compile(r"\[(.*?)\]\((.*?)\)")
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)")
CODE_RE = re.compile(r"`([^`]+)`")


def _safe_text(value: str) -> str:
    text = str(value or "")
    text = LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)
    text = BOLD_RE.sub(lambda m: m.group(1), text)
    text = ITALIC_RE.sub(lambda m: m.group(1), text)
    text = CODE_RE.sub(lambda m: m.group(1), text)
    return text.strip()


def _image_match(line: str):
    return IMAGE_RE.fullmatch(line.strip())


def _fetch_image_bytes(url: str) -> bytes | None:
    if not url:
        return None
    try:
        response = requests.get(url, timeout=20)
        response.raise_for_status()
        content_type = (response.headers.get("content-type") or "").lower()
        if "image" not in content_type and not re.search(r"\.(png|jpe?g|gif|webp|bmp)(\?|$)", url, flags=re.I):
            return None
        return response.content
    except Exception:
        return None


def _display_doc_type(document: Dict[str, Any]) -> str:
    value = str(document.get("doc_type") or "doc").strip().replace("_", " ")
    return value.upper() if value else "DOC"


def _base_filename(document: Dict[str, Any]) -> str:
    title = str(document.get("title") or "Document").strip()
    title = re.sub(r'[\\/:*?"<>|]+', " ", title)
    title = re.sub(r"\s+", " ", title).strip() or "Document"
    return f"{title} - {_display_doc_type(document)}"


def _source_label(document: Dict[str, Any]) -> str:
    return str(document.get("source_label") or "Internal workflow")


def build_docx_bytes(document: Dict[str, Any]) -> Tuple[bytes, str]:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title = str(document.get("title") or "Untitled Document")
    summary = str(document.get("summary") or "").strip()
    content = str(document.get("content") or "")

    heading = doc.add_heading(title, 0)
    heading.runs[0].font.size = Pt(22)

    meta = []
    if document.get("doc_type"):
        meta.append(_display_doc_type(document))
    if document.get("source_label"):
        meta.append(_source_label(document))
    if document.get("created_at"):
        meta.append(str(document.get("created_at")))
    if meta:
        p = doc.add_paragraph(" • ".join(meta))
        p.runs[0].italic = True

    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        image = _image_match(stripped)
        if image:
            alt_text, url = image.group(1), image.group(2)
            image_bytes = _fetch_image_bytes(url)
            if image_bytes:
                doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.0))
                if alt_text:
                    cap = doc.add_paragraph(alt_text)
                    cap.runs[0].italic = True
            else:
                doc.add_paragraph(_safe_text(stripped))
            continue

        if stripped.startswith("#"):
            level = max(1, min(3, len(stripped) - len(stripped.lstrip("#"))))
            doc.add_heading(_safe_text(stripped.lstrip("#").strip()), level=level)
            continue

        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(_safe_text(stripped[2:].strip()), style="List Bullet")
            continue

        doc.add_paragraph(_safe_text(stripped))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue(), f"{_base_filename(document)}.docx"


def build_pdf_bytes(document: Dict[str, Any]) -> Tuple[bytes, str]:
    title = str(document.get("title") or "Untitled Document")
    summary = str(document.get("summary") or "").strip()
    content = str(document.get("content") or "")

    output = io.BytesIO()
    pdf = SimpleDocTemplate(
        output,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]
    body.spaceAfter = 8
    bullet_style = ParagraphStyle("BulletBody", parent=body, leftIndent=12, bulletIndent=0, spaceAfter=6)
    meta_style = ParagraphStyle("Meta", parent=body, textColor="#625f6b", fontSize=10, italic=True, spaceAfter=10)

    story = [Paragraph(_safe_text(title), title_style)]

    meta = []
    if document.get("doc_type"):
        meta.append(_display_doc_type(document))
    if document.get("source_label"):
        meta.append(_source_label(document))
    if document.get("created_at"):
        meta.append(str(document.get("created_at")))
    if meta:
        story.append(Paragraph(_safe_text(" • ".join(meta)), meta_style))

    if summary:
        story.append(Paragraph("Summary", h1))
        story.append(Paragraph(_safe_text(summary), body))
        story.append(Spacer(1, 0.08 * inch))

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        image = _image_match(stripped)
        if image:
            alt_text, url = image.group(1), image.group(2)
            image_bytes = _fetch_image_bytes(url)
            if image_bytes:
                img = RLImage(io.BytesIO(image_bytes))
                max_width = 6.4 * inch
                if img.drawWidth > max_width:
                    ratio = max_width / float(img.drawWidth)
                    img.drawWidth = max_width
                    img.drawHeight = float(img.drawHeight) * ratio
                story.append(img)
                if alt_text:
                    story.append(Paragraph(_safe_text(alt_text), meta_style))
            else:
                story.append(Paragraph(_safe_text(stripped), body))
            continue

        if stripped.startswith("#"):
            level = max(1, min(2, len(stripped) - len(stripped.lstrip("#"))))
            style = h1 if level == 1 else h2
            story.append(Paragraph(_safe_text(stripped.lstrip("#").strip()), style))
            continue

        if stripped.startswith(("- ", "* ")):
            story.append(Paragraph(f"• {_safe_text(stripped[2:].strip())}", bullet_style))
            continue

        story.append(Paragraph(_safe_text(stripped), body))

    pdf.build(story)
    return output.getvalue(), f"{_base_filename(document)}.pdf"


def build_email_draft_bytes(document: Dict[str, Any], attachment_format: str = "docx") -> Tuple[bytes, str]:
    attachment_format = (attachment_format or "docx").lower()
    if attachment_format == "pdf":
        attachment_bytes, attachment_name = build_pdf_bytes(document)
        maintype, subtype = "application", "pdf"
    else:
        attachment_bytes, attachment_name = build_docx_bytes(document)
        maintype, subtype = "application", "vnd.openxmlformats-officedocument.wordprocessingml.document"

    title = str(document.get("title") or "Document")
    summary = str(document.get("summary") or "").strip()

    msg = EmailMessage(policy=SMTP)
    msg["Subject"] = f"Process document: {title}"
    msg["To"] = ""
    msg["Cc"] = ""
    msg["Bcc"] = ""

    body_lines = [
        "Hi,",
        "",
        f"Attached is the requested process document: {title}.",
    ]
    if summary:
        body_lines.extend(["", "Summary:", summary])
    body_lines.extend(["", "Thanks,"])

    msg.set_content("\n".join(body_lines))
    msg.add_attachment(attachment_bytes, maintype=maintype, subtype=subtype, filename=attachment_name)

    filename = f"{_base_filename(document)} - Email Draft.eml"
    return msg.as_bytes(), filename